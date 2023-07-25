from PySide2.QtCore import (QCoreApplication, QMetaObject, QObject, QPoint,
    QRect, QSize, QUrl, Qt, QFile, QTextStream, QIODevice, Signal, Qt)
from PySide2.QtGui import (QBrush, QColor, QConicalGradient, QCursor, QFont,
    QFontDatabase, QIcon, QLinearGradient, QPalette, QPainter, QPixmap,
    QRadialGradient, QResizeEvent)
from PySide2.QtWidgets import *
import sys, os, re
import constants as gc
import helpers as gf
import sys
import os
from docx import Document
import openpyxl
import win32com.client
from math import modf
from docx2pdf import convert

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS2
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

actualPage = list(gc.PAGES_TITLES.keys())[0]

oShell = win32com.client.Dispatch("Wscript.Shell")
pathToMyDocuments = oShell.SpecialFolders("MyDocuments")

class ClassWindow(QMainWindow):
    resized = Signal( QResizeEvent )
    def __init__(self, parent = None) -> None:
        super().__init__(parent)
        if self.objectName():
            self.setObjectName(list(gc.PAGES_TITLES.values())[0])
        self.resize(678, 497)
        icon = QIcon()
        icon.addFile(u":/newPrefix/Message To Document 2.png", QSize(), QIcon.Normal, QIcon.Off)
        self.setWindowIcon(icon)
        self.centralwidget = QWidget(self)
        self.centralwidget.setObjectName(u"centralwidget")
        self.horizontalLayout = QHBoxLayout(self.centralwidget)
        self.horizontalLayout.setSpacing(0)
        self.horizontalLayout.setObjectName(u"horizontalLayout")
        self.horizontalLayout.setContentsMargins(0, 0, 0, 0)
        self.stackedWidget = QStackedWidget(self.centralwidget)
        self.stackedWidget.setObjectName(u"stackedWidget")

        self.horizontalLayout.addWidget(self.stackedWidget)

        self.setCentralWidget(self.centralwidget)

        self.retranslateUi()

        QMetaObject.connectSlotsByName(self)
        
    def resizeEvent(self, event: QResizeEvent) -> None:
        self.resized.emit(event)
        return super(ClassWindow, self).resizeEvent(event)

    def retranslateUi(self):
        self.setWindowTitle(QCoreApplication.translate("Message To Doc", list(gc.PAGES_TITLES.values())[0], None))
    # retranslateUi
    

if __name__ == "__main__":
    styleSheetFile = QFile( resource_path("style.qss") )
    styleSheetFile.open(QIODevice.ReadOnly)
    styleSheet = QTextStream(styleSheetFile).readAll()
    form = QApplication(sys.argv)
    form.setStyleSheet(styleSheet)
    fontDatabase = QFontDatabase()
    fontDatabase.addApplicationFont(resource_path("fonts\\inter.ttf"))
    form.setFont(QFont(fontDatabase.applicationFontFamilies(0)[0], 10))
    window = ClassWindow() #gf.load_ui("window")
    #window.setWindowTitle( QCoreApplication.translate("Template", "Template", None) )
    
    def activePage(page: str) :
        for _page in gc.PAGES.keys() :
            if page == _page :
                window.__getattribute__(page).m_ui.__getattribute__(_page).setStyleSheet(headerButtonActiveStyle)
            else :
                window.__getattribute__(page).m_ui.__getattribute__(_page).setStyleSheet(headerButtonStyle)
                
    def compile_to_word():
        """Compile the QPlainTextEdit's text to a Microsoft Word Document
        """        
        title_doc = window.word.m_ui.titleEdit.text() if window.word.m_ui.titleEdit.text()!='' else gc.DEFAULT_TITLES_DOC["word"]
        doc = Document(resource_path("template.docx"))
        doc.add_heading(title_doc, 0)
        doc.add_page_break()
        text_treated = gf.treat_text(window.word.m_ui.messageEdit.toPlainText())
        #print(text_treated)
        for i in text_treated.keys():
            doc.add_heading(text_treated[i]["title"] , 1 if modf(float(i))[0] == 0 else 2 )
            for j in text_treated[i]["content"] :
                if text_treated[i]["content"][j]["type"] == "normal" :
                    paragraph = doc.add_paragraph(text_treated[i]["content"][j]["content"])
                else :
                    for k in range(len(text_treated[i]["content"][j]["content"])) :
                        paragraph = doc.add_paragraph(text_treated[i]["content"][j]["content"][k], style = text_treated[i]["content"][j]["type"])
        #print(QFileDialog.getSaveFileName(None, gc.FILE_DIALOG_CAPTION, "", "Microsoft Word Files (*.docx)"))
        pathWhereSave = QFileDialog.getSaveFileName(None, gc.FILE_DIALOG_CAPTION, os.path.join( pathToMyDocuments, f"{title_doc}.docx" ), "Microsoft Word Files (*.docx)")
        if pathWhereSave[0] != '' :
            doc.save(pathWhereSave[0])
        
    def compile_to_excel():
        title_doc = window.excel.m_ui.titleEdit.text() if window.excel.m_ui.titleEdit.text()!='' else gc.DEFAULT_TITLES_DOC["excel"]
        shouldTotal = window.excel.m_ui.totalColumn.isChecked()
        docExcel = openpyxl.Workbook()
        doc = docExcel.worksheets[0]
        text = gf.treat_text_excel(window.excel.m_ui.messageEdit.toPlainText(), delimiter= window.excel.m_ui.delimiterEdit.text(), shouldTotal= shouldTotal) 
        for row in text :
            doc.append(row)
        pathWhereSave = QFileDialog.getSaveFileName(None, gc.FILE_DIALOG_CAPTION, os.path.join( pathToMyDocuments, f"{title_doc}.xlsx" ), "Microsoft Excel Files (*.xlsx *.xls)")
        if pathWhereSave[0] != '' :
            docExcel.save(pathWhereSave[0])
        
    def compile_to_pdf():
        title_doc = window.pdf.m_ui.titleEdit.text() if window.pdf.m_ui.titleEdit.text()!='' else gc.DEFAULT_TITLES_DOC["pdf"]
        shouldConvert = window.pdf.m_ui.wordConvert.isChecked()
        pathWhereSave = QFileDialog.getSaveFileName(None, gc.FILE_DIALOG_CAPTION, os.path.join( pathToMyDocuments, f"{title_doc}.pdf" ), "PDF Files (*.pdf)")
        if pathWhereSave[0] != '' :
            if shouldConvert :
                input_file = gf.compile_to_word(window.pdf.m_ui.messageEdit.toPlainText())
                convert(input_file, pathWhereSave[0])
            else :
                gf.text_to_pdf(window.pdf.m_ui.messageEdit.toPlainText(), pathWhereSave[0])
            
    def importTxtFile(page) :
        pathToExport = QFileDialog.getOpenFileName(None, gc.IMPORT_FILE_DIALOG_CAPTION, pathToMyDocuments, "TXT Files (*.txt *.csv)")
        if pathToExport[0] != 0 :
            with open(pathToExport[0]) as f :
                window.__getattribute__(page).m_ui.messageEdit.setPlainText(f.read())
    
    def goToPage(page : str) :
        """Go to the specified page

        Args:
            page (str): The destination page
        """        
        actualPage = page
        window.setWindowTitle( QCoreApplication.translate("Message To Doc", gc.PAGES_TITLES[page], None) )
        window.stackedWidget.setCurrentIndex(gc.PAGES_INDEX[page])
        activePage(page)
    
    # Create pages and load ui files in it
    window.__setattr__("word", gf.load_py("word"))
    window.stackedWidget.addWidget(window.word)
    window.word.m_ui.word.clicked.connect(lambda : goToPage("word"))
    window.word.m_ui.excel.clicked.connect(lambda : goToPage("excel"))
    window.word.m_ui.pdf.clicked.connect(lambda : goToPage("pdf"))
    window.word.m_ui.compile.clicked.connect(lambda : compile_to_word())
    window.__setattr__("excel", gf.load_py("excel"))
    window.stackedWidget.addWidget(window.excel)
    window.excel.m_ui.word.clicked.connect(lambda : goToPage("word"))
    window.excel.m_ui.excel.clicked.connect(lambda : goToPage("excel"))
    window.excel.m_ui.pdf.clicked.connect(lambda : goToPage("pdf"))
    window.excel.m_ui.compile.clicked.connect(lambda : compile_to_excel())
    window.__setattr__("pdf", gf.load_py("pdf"))
    window.stackedWidget.addWidget(window.pdf)
    window.pdf.m_ui.word.clicked.connect(lambda : goToPage("word"))
    window.pdf.m_ui.excel.clicked.connect(lambda : goToPage("excel"))
    window.pdf.m_ui.pdf.clicked.connect(lambda : goToPage("pdf"))
    window.pdf.m_ui.compile.clicked.connect(lambda : compile_to_pdf())
    for page in gc.PAGES.keys():
        window.__getattribute__(page).m_ui.takeFileTxt.clicked.connect(lambda : importTxtFile(page))
    
    # Retrieve and change the styles of header's buttons
    headerButtonStyle = window.pdf.m_ui.pdf.styleSheet()
    headerButtonActiveStyle = headerButtonStyle + "color: rgb(79, 149, 190);"
    window.word.m_ui.word.setStyleSheet(headerButtonActiveStyle)
    
    window.show()
    form.exec_()