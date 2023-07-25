import sys
import os
import re
from PySide2.QtWidgets import QWidget
from PySide2.QtCore import (QFile, QIODevice)
from PySide2.QtUiTools import QUiLoader
import constants as gc
import textwrap
from fpdf import FPDF
import math
from docx import Document

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS2
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)
    
def load_ui(file_name:str) -> QWidget:
    """Load an ui file in a QWidget and return it

    Args:
        file_name (str): The name of the specified ui file

    Returns:
        QWidget: The resulted QWidget
    """    
    ui_page_file = QFile(gc.PAGES[file_name])
    if not ui_page_file.open(QIODevice.ReadOnly):
        print(f"Cannot open {file_name} : {ui_page_file.errorString()}")
        sys.exit(-1)
    loader = QUiLoader()
    window = loader.load(ui_page_file)
    ui_page_file.close()
    if not window:
        print(loader.errorString())
        sys.exit(-1)
    return window  

def load_py(file_name:str) -> QWidget:
    """Load python compiled ui file in QWidget

    Args:
        file_name (str): The name of the compiled python file from ui file

    Returns:
        QWidget: The resulted QWidget
    """    
    class _Window(QWidget):
        def __init__(self, parent=None):
            super(_Window, self).__init__(parent)

            self.m_ui = gc.PAGES_UI[file_name]
            self.m_ui.setupUi(self)
    return _Window()

def escape_before(_str_to_escape: str, _str: str) -> str :
    """Get the characters behind specified characters within a string

    Args:
        _str_to_escape (str): The string placed before the researched characters
        _str (str): The string in which we have to search

    Returns:
        str: The researched characters
    """        
    pattern_string = re.escape(_str_to_escape) + "(.*)"
    pattern = re.compile(pattern_string)
    return pattern.findall(_str)[0]

def escape_before_with_pattern(_pattern_to_escape: str, _str: str) -> str :
    """Get the characters behind a specific pattern within a string

    Args:
        _pattern_to_escape (str): The pattern to escape
        _str (str): The string to search in

    Returns:
        str: The researched characters
    """    
    _str_to_escape = re.findall( _pattern_to_escape, _str )
    return escape_before(_str_to_escape[0] if len(_str_to_escape) != 0 else '' , _str)

def treat_text(_text: str) -> dict[float:dict[str:str]] :
    """Treat a text and return a dictionnary of paragraphs like that 
        {
            index[float] : {
                "title" : str,
                "content" : {
                    index[int] : {
                        "type" : str,
                        "content" : str|list[str],
                    },
                    ...
                }
            },
            ...
        }

    Args:
        _text (str): a multiline string

    Returns:
        dict[float:dict[str:str]]: dictionnary of paragraphs
    """    
    text = _text.splitlines(True)
    text = [ i for i in text if i!='\n' ]
    text = order_text(text)
    
    # Create a list of all index in text and order sub-index
    list_index = []
    actual_higher_level = 0
    #print(text_treated)
    for _paragraph in text:
        
        left_delimiter = re.findall(r'\d{1,3}\. ', _paragraph) if len(re.findall(r'\d{1,3}\. ', _paragraph)) > 0 else re.findall(r'\d{1,3}\.', _paragraph)
        left_delimiter_2 = re.findall(r' \d{1,3}\. ', _paragraph) if len(re.findall(r' \d{1,3}\. ', _paragraph)) > 0 else re.findall(r' \d{1,3}\.', _paragraph)
        #if len(left_delimiter) > 0 or len(left_delimiter_2) > 0 :
        left_delimiter_int = int(re.findall(r'\d', left_delimiter[0])[0]) if len(left_delimiter) > 0 else 0
        left_delimiter_int_2 = int(re.findall(r'\d', left_delimiter_2[0])[0]) if len(left_delimiter_2) != 0 else 0
        actual_higher_level = ( left_delimiter_int if left_delimiter_int != 0 else actual_higher_level ) if left_delimiter_int_2 == 0 else actual_higher_level
        list_index.append( actual_higher_level if left_delimiter_int_2 == 0 else concatNumbersToFloat(actual_higher_level, left_delimiter_int_2)  )
    
    # Create an object for the text and fill it along the list of index
    text_treated_2 = {}
    index_in_list = 0
    for _paragraph in text:
        
        left_delimiter = re.findall(r'\d{1,3}\. ', _paragraph)
        left_delimiter_2 = re.findall(r'\d{1,3}\.', _paragraph)
        #left_delimiter_int = int(re.findall(r'\d', left_delimiter[0])[0])
        #left_delimiter_int_2 = int(re.findall(r'\d', left_delimiter_2[0])[0])
        right_delimiter = re.findall( r':([\s\S]*)' ,_paragraph)
        right_delimiter_2 = re.findall( r'\n([\s\S]*)' ,_paragraph)
        
        left_string = left_delimiter[0] if len(left_delimiter) != 0 else left_delimiter_2[0] if len(left_delimiter_2) != 0 else ''
        pattern_string = ( re.escape(left_string) ) + "(.*:)"
        pattern_string_2 = ( re.escape(left_string) ) + "(.*)"
        pattern = re.compile(pattern_string)
        pattern_2 = re.compile(pattern_string_2)
        title = pattern.findall(_paragraph)[0] if len(pattern.findall(_paragraph)) > 0 else pattern_2.findall(_paragraph)[0]
        
        splitted_paragraph = right_delimiter[0].splitlines(True) if len(right_delimiter) > 0 else right_delimiter_2[0].splitlines(True)
        paragraph_treated = {}
        actual_paragraph_position = 0
        for i in range(len(splitted_paragraph)) :
            point_character = re.findall(r'^\s*•', splitted_paragraph[i])
            hyphen_character = re.findall(r'^\s*-', splitted_paragraph[i])
            if len(point_character) != 0 :
                if i > 0 and len(re.findall(r'^\s*•', splitted_paragraph[i-1])) != 0 :
                    paragraph_treated[actual_paragraph_position]["content"].append(escape_before_with_pattern(r'^\s*•', splitted_paragraph[i]))
                else :
                    actual_paragraph_position += 1 if i!=0 else 0
                    paragraph_treated[actual_paragraph_position] = {
                        "type" : "Point",
                        "content" : [escape_before_with_pattern(r'^\s*•', splitted_paragraph[i])]
                    }
            elif len(hyphen_character) != 0 :
                if i > 0 and len(re.findall(r'^\s*-', splitted_paragraph[i-1])) != 0 :
                    paragraph_treated[actual_paragraph_position]["content"].append(escape_before_with_pattern(r'^\s*-', splitted_paragraph[i]))
                else :
                    actual_paragraph_position += 1 if i!=0 else 0
                    paragraph_treated[actual_paragraph_position] = {
                        "type" : "Hyphen",
                        "content" : [escape_before_with_pattern(r'^\s*-', splitted_paragraph[i])]
                    }
            else :
                actual_paragraph_position += 1 if i!=0 else 0
                paragraph_treated[actual_paragraph_position] = {
                    "type" : "normal",
                    "content" : splitted_paragraph[i]
                }
        
        text_treated_2[list_index[index_in_list]] = {
            "title" : title,
            "content" : paragraph_treated
        }
        index_in_list += 1

    return text_treated_2

def order_text(_text: list[str], i: int = 0) -> list[str]:
    """Order a text by appending to the end of each title the content of paragraph so that we can obtain an ordered list of paragraphs

    Args:
        _text (list[str]): The specified and already splitted text to order
        i (int, optional): The actual index of text. Defaults to 0.

    Returns:
        list[str]: A list of ordered paragraphs of the splitted text
    """    
    if not bool(re.search(r'\d\.', _text[i])) :
        if i != 0 and bool(re.search(r'\d\.', _text[i-1])) :
            _text[i-1] += _text[i] # f"""{_text[i-1]} \n {_text[i]} """
            _text[i] = ''
            _text = [ i for i in _text if i!='' ]
        else :
            i+=1
    else :
        i+=1
    return _text if i == len(_text) else order_text(_text, i)        

def  take_title(_paragraph: str) :
    """Take the title of paragraph or string

    Args:
        _string (str): A paragraph or string

    Returns:
        str, str: The title of the paragraph or string, The content of the paragraph
    """    
    left_delimiter = re.findall(r'^\d{1,3}\. ', _paragraph)
    right_delimiter = re.findall(r':(.*)$', _paragraph)
    pattern_string = ( re.escape(left_delimiter[0]) if len(right_delimiter) != 0 else '' ) + "(.*)" + ( re.escape(right_delimiter[0]) if len(right_delimiter) != 0 else '' )
    pattern = re.compile(pattern_string)
    return pattern.findall(_paragraph)[0], right_delimiter

def input_to_format_date(formatInput: str) -> str:
    """Format a AAAA/MM/JJ to proper date format

    Args:
        formatInput (str): The value of the user input

    Returns:
        str: The resulted proper date format
    """    
    formatInput.replace("AAAA","%Y")
    formatInput.replace("MM","%m")
    formatInput.replace("JJ","%d")

def treat_text_excel(_text: str, delimiter: str = ',', shouldTotal: bool = False) -> list[list[str]] :
    """Give an array of each line of a text, where each line are them-selves splitted by a delimiter character and optionally add a total value for each line

    Args:
        _text (str): The specified text to double-splt
        delimiter (str, optional): The delimiter character to split each line of the text. Defaults to ','.
        shouldTotal(bool): A bool who determines if the total column should be add. Defaults to False.

    Returns:
        list[list[str]]: The resulted array of lines which are them-selves array of string
    """    
    delimiter = ',' if delimiter == '' else delimiter
    text = [ i.split(delimiter) for i in _text.split('\n') ]
    if shouldTotal :
        for i in range(len(text)) :
            intValues = [int(x) for x in text[i] if (x.isdigit() or x[0] == '-' and x[1:].isdigit())]
            total = sum(intValues)
            if len(intValues) != 0 :
                text[i].append(total)
    return text

def text_to_pdf(text: str, filename: str):
    a4_width_mm = 210
    pt_to_mm = 0.35
    fontsize_pt = 10
    fontsize_mm = fontsize_pt * pt_to_mm
    margin_bottom_mm = 10
    character_width_mm = 7 * pt_to_mm
    width_text = a4_width_mm / character_width_mm

    pdf = FPDF(orientation='P', unit='mm', format='A4')
    pdf.set_auto_page_break(True, margin=margin_bottom_mm)
    pdf.add_page()
    pdf.set_font(family='Courier', size=fontsize_pt)
    splitted = [ str(i.encode("latin-1", "replace").decode("latin-1")) for i in text.split('\n')]

    for line in splitted:
        lines = textwrap.wrap(line, width_text)

        if len(lines) == 0:
            pdf.ln()

        for wrap in lines:
            pdf.cell(0, fontsize_mm, wrap, ln=1)

    pdf.output(filename, 'F')

def compile_to_word(_text: str) -> str :
    """Compile text to word and return the path of the temporary file

    Args:
        _text (str): The text to compile

    Returns:
        str: The path of the temporary file
    """    
    title_doc = gc.DEFAULT_TITLES_DOC["word"]
    doc = Document(resource_path("template.docx"))
    doc.add_heading(title_doc, 0)
    doc.add_page_break()
    text_treated = treat_text(_text)
    #print(text_treated)
    for i in text_treated.keys():
        doc.add_heading(text_treated[i]["title"] , 1 if math.modf(float(i))[0] == 0 else 2 )
        for j in text_treated[i]["content"] :
            if text_treated[i]["content"][j]["type"] == "normal" :
                paragraph = doc.add_paragraph(text_treated[i]["content"][j]["content"])
            else :
                for k in range(len(text_treated[i]["content"][j]["content"])) :
                    paragraph = doc.add_paragraph(text_treated[i]["content"][j]["content"][k], style = text_treated[i]["content"][j]["type"])
    #print(QFileDialog.getSaveFileName(None, gc.FILE_DIALOG_CAPTION, "", "Microsoft Word Files (*.docx)"))
    pathTempSaved = resource_path(f"temporary_{title_doc}.docx")
    doc.save(pathTempSaved)
    return pathTempSaved

  
def countDigit(n):
    return math.floor(math.log10(n)+1)-1

def concatNumbersToFloat(first: int, second: int) -> float:
    """Concatenate two integers a and b to float a,b 

    Args:
        first (int): The integer before the comma
        second (int): The integer after the comma

    Returns:
        float: The resulted float concatenated
    """
    return first+(second/(10*pow(10,(countDigit(second)))))