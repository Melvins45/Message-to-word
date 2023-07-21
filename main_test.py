from docx import Document
import re

"""document = Document()
document.add_heading('A simple text', level=1)
document.add_paragraph('some more text ... ')
paragraph = document.add_paragraph('yes')
paragraph.style = document.styles['Heading 4']"""

"""document = Document("template.docx")
paragraph = document.add_paragraph('yes')
print(document.styles.keys())
paragraph.style = document.styles['Numbering1']
document.add_heading('My first numbered heading', 1)
document.add_heading('My first numbered heading 2', 1)

document.save('docx_file.docx')"""

text = {
    'Analyse de besoins' : """Pour faire une analyse des besoins d'un projet de site web, il est important de suivre les étapes suivantes :

1. Identifier les objectifs du site web : Il est important de définir les objectifs du site web, tels que la vente de produits, la promotion d'une entreprise ou d'un service, la collecte de données, etc.

2. Identifier les publics cibles : Il est important de déterminer les différents publics cibles du site web, tels que les clients potentiels, les employés, les partenaires commerciaux, etc.

3. Analyser les besoins des utilisateurs : Il est important de comprendre les besoins des utilisateurs en termes de fonctionnalités, d'interface utilisateur et de contenu.

4. Analyser les besoins techniques : Il est important de comprendre les exigences techniques pour le développement du site web, telles que la plate-forme de développement, le langage de programmation, la base de données, etc.

5. Identifier les contraintes et les ressources disponibles : Il est important de comprendre les contraintes du projet, telles que le budget, le temps et les ressources disponibles.

6. Prioriser les besoins : Il est important de hiérarchiser les besoins en fonction de leur importance pour le succès du projet.

7. Réaliser une étude de marché : Il est important de réaliser une étude de marché pour comprendre les tendances actuelles dans l'industrie et identifier les concurrents.

8. Élaborer un plan d'action : Une fois que tous les besoins ont été identifiés et analysés, il est important d'élaborer un plan d'action pour répondre à ces besoins et atteindre les objectifs du projet.""",

    'Plannification' : """Pour faire la planification d'un projet de site web, il est important de suivre les étapes suivantes :

1. Définir les objectifs du projet : Il est important de définir les objectifs du projet, tels que la création d'un nouveau site web, la refonte d'un site existant, l'ajout de nouvelles fonctionnalités, etc.

2. Identifier les parties prenantes : Il est important de déterminer les parties prenantes du projet, telles que les clients, les utilisateurs finaux, les développeurs, les designers, etc.

3. Élaborer un budget : Il est important d'élaborer un budget pour le projet, en prenant en compte les coûts de développement, de conception, d'hébergement et de maintenance.

4. Élaborer un calendrier : Il est important d'établir un calendrier réaliste pour le projet, en prenant en compte les délais de développement, de conception et de test.

5. Identifier les ressources nécessaires : Il est important d'identifier les ressources nécessaires pour le projet, telles que les compétences techniques, le matériel informatique et les logiciels.

6. Élaborer un plan de communication : Il est important d'élaborer un plan de communication pour le projet, en prenant en compte les besoins de communication entre les parties prenantes.

7. Élaborer un plan de gestion des risques : Il est important d'élaborer un plan de gestion des risques pour le projet, en identifiant les risques potentiels et en élaborant des stratégies pour les gérer.

8. Élaborer un plan de test et de validation : Il est important d'élaborer un plan de test et de validation pour le projet, en prenant en compte les exigences de qualité et de sécurité.

9. Élaborer un plan de maintenance : Il est important d'élaborer un plan de maintenance pour le projet, en prenant en compte les besoins de mise à jour et de correction des erreurs.

10. Élaborer un plan de formation : Il est important d'élaborer un plan de formation pour le projet, en prenant en compte les besoins de formation des utilisateurs finaux et des parties prenantes.""",

    'Conception' : """Pour concevoir un projet de site web, il est important de suivre les étapes suivantes :

1. Analyser les besoins : Il est important d'analyser les besoins des utilisateurs finaux et des parties prenantes pour déterminer les fonctionnalités et les exigences du site web.

2. Définir l'architecture du site : Il est important de définir l'architecture du site, en créant une structure de navigation claire et en déterminant les pages et les sections du site.

3. Créer un wireframe : Il est important de créer un wireframe, qui est une représentation visuelle de la disposition des éléments sur chaque page du site.

4. Concevoir la maquette graphique : Il est important de concevoir la maquette graphique du site, en créant un design attractif et cohérent avec l'image de marque de l'entreprise.

5. Intégrer les fonctionnalités : Il est important d'intégrer les fonctionnalités nécessaires, telles que les formulaires de contact, les boutons de partage sur les réseaux sociaux, les options de paiement en ligne, etc.

6. Optimiser le contenu : Il est important d'optimiser le contenu du site pour améliorer le référencement naturel et l'expérience utilisateur.

7. Tester le site : Il est important de tester le site pour s'assurer qu'il fonctionne correctement et qu'il répond aux exigences de qualité et de sécurité.

8. Lancer le site : Il est important de lancer le site et de s'assurer qu'il est accessible à tous les utilisateurs.

9. Assurer la maintenance : Il est important d'assurer la maintenance du site pour corriger les erreurs et mettre à jour le contenu et les fonctionnalités.

10. Mesurer les résultats : Il est important de mesurer les résultats du site pour évaluer son efficacité et déterminer les améliorations à apporter."""
}

def treat_text(_text: str) :
    """Treat a text and return a list of paragraphs

    Args:
        _text (str): a multiline string

    Returns:
        list[str]: list of paragraphs
    """    
    text = _text.split('\n')
    return [ i for i in text if i!='' and bool(re.search(r'\d', i)) ]

# Split the message in list of paragraphs
"""
test = start_string.split('\n')

# Remove empty strings and non-titled paragraphs
test = [ i for i in test if i!='' and bool(re.search(r'\d', i)) ]"""

def  take_title(_paragraph: str) :
    """Take the title of paragraph or string

    Args:
        _string (str): A paragraph or string

    Returns:
        str: The title of the paragraph or string
    """    
    left_delimiter = re.findall(r'^\d{1,3}\. ', _paragraph)
    right_delimiter = re.findall(r':(.*)$', _paragraph)
    pattern_string = re.escape(left_delimiter[0]) + "(.*)" + re.escape(right_delimiter[0])
    pattern = re.compile(pattern_string)
    return pattern.findall(_paragraph)[0]

#testy = pattern.findall(test[0])
#test = [ take_title(i) for i in test ]

#print(test)
#print(testy, test[0], pattern_string)

doc = Document("template.docx")
doc.add_heading('Réalisation du projet de site web de MTW', 0)
doc.add_page_break()
print([i for i in text])

for first_level_heading in text.keys() :
    doc.add_heading(first_level_heading, 1)
    for second_level_heading in treat_text(text[first_level_heading]) :
        doc.add_heading(take_title(second_level_heading), 2)
        doc.add_paragraph('\n\n\n')
    
doc.save('Message To Word.docx')